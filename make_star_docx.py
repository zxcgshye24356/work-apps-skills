# -*- coding: utf-8 -*-
"""生成 STAR 作品 Word 文档（百度收录检测工具），含 6 张截图。"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

CJK = "Microsoft YaHei"
BASE = r"C:\Users\liu\WorkBuddy\2026-08-07-16-30-00\baidu_index_checker"
IMG = lambda n: f"{BASE}\\screenshots\\{n}.png"

doc = Document()

# 基础样式（中文友好）
normal = doc.styles["Normal"]
normal.font.name = CJK
normal.font.size = Pt(11)
normal.element.rPr.get_or_add_rFonts().set(qn("w:eastAsia"), CJK)


def set_cjk(run, name=CJK):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    return run


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    set_cjk(r)
    return p


def add_sub(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_cjk(r)
    return p


def add_h(text):
    p = doc.add_heading(level=1)
    set_cjk(p.add_run(text))
    return p


def add_para(text):
    p = doc.add_paragraph()
    set_cjk(p.add_run(text))
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        set_cjk(p.add_run(bold_prefix)).bold = True
    set_cjk(p.add_run(text))
    return p


def add_runs(p, parts):
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        set_cjk(r)


def add_image(num, caption):
    doc.add_picture(IMG(num), width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    set_cjk(r)


# ---- 标题 ----
add_title("百度收录效果批量检测工具（双品牌词版）")
add_sub("项目作品 · 用于求职展示  |  角色：产品需求方 + AI 协作开发（非技术背景，借助 AI 编程工具完成从 0 到 1 的搭建与迭代）")

# ---- S ----
add_h("一、项目背景（Situation）")
add_para("我在一家化工 / 生物类产品内容公司做产品内容。我们公司手上有两个品牌：「陕西新研博美」和「西安凯新生物」。同一个产品（比如\"聚乙二醇\"\"DSPE-PEG\"\"巯基聚乙二醇\"）往往两个品牌下都在卖。")
add_para("做内容运营时有个很现实的问题：用户在百度搜一个产品时，到底是哪个品牌的内容被百度收录得更多、曝光更好？这直接决定我们该用哪个品牌账号继续写文章、把精力投在哪边。")
add_para("以前这件事是纯人工做的：打开百度 → 搜产品名 → 一页页点开文章 → 肉眼看哪篇提到了\"陕西新研博美\"、哪篇提到了\"西安凯新生物\" → 手动数、手动记。几十个产品一轮下来又慢又容易数错，而且百度时不时弹验证码，人更崩溃。")

# ---- T ----
add_h("二、我的任务（Task）")
add_para("我想要一个能批量处理的工具：")
add_bullet("我只要上传一张产品清单（Excel）；")
add_bullet("它自动去百度第一页搜每个产品；")
add_bullet("自动把搜出来的文章一篇篇打开，判断正文里有没有出现两个品牌词；")
add_bullet("最后给我一份带\"归到哪个品牌 / 各品牌提了多少篇 / 收录效果好不好\"的 Excel 报告。")
add_para("核心难点有两个：一是要绕开百度的反爬（不能简单爬，得模拟真人浏览器）；二是\"一篇文章只算一次品牌词\"这种计数规则必须准，否则结论就歪了。")

# ---- A ----
add_h("三、我做的事（Action）")
add_para("作为一个不写代码的人，我把自己当成\"产品 + 项目经理\"，用 AI 编程工具把需求一步步落成了一个能用的 Web 应用。技术栈是 Python + Streamlit + Playwright + pandas / openpyxl。")

add_h("1. 先把需求想清楚，再交给 AI 实现")
add_para("我没有上来就让 AI 写代码，而是先借另一个 AI（DeepSeek）把模糊的想法梳理成结构化的需求文档，自己审一遍把漏洞补上，再让 AI 输出成 md 格式的需求书，确认无误后才发给 WorkBuddy 落地实现。")
add_image(1, "图 1：先用 AI 把\"百度收录对比\"这个模糊想法，梳理成有背景、输入、处理逻辑、输出的结构化需求。")
add_image(2, "图 2：自己审一遍，把需求里没考虑到的边界情况（重复产品、搜不到、链接打不开等）反馈补充进去。")
add_image(3, "图 3：确认后的需求文档（md 格式），作为交给实现端（WorkBuddy）的\"说明书\"。")

add_h("2. 第一版：先把核心功能跑通")
add_bullet("上传 Excel 自动去重，重复的会提示出来、不重复搜；", "去重：")
add_bullet("用 Playwright 驱动真实 Chromium 浏览器去百度搜，带浏览器标识、抹掉自动化痕迹，规避反爬；", "反爬：")
add_bullet("抓取百度第一页的文章链接，逐篇打开提取正文；", "抓取：")
add_bullet("判断每篇文章是否提到两个品牌词，严格执行\"一篇文章对每个品牌只计一次\"——即使一篇同时提到两个品牌也不重复计数；", "计数：")
add_bullet("比较两个品牌被提到的篇数，多的归它，一样多标\"持平\"；", "归类：")
add_bullet("总提及 ≥ 5 篇标\"好\"，否则标\"差\"；", "判定：")
add_bullet("输出 6 列 Excel：产品名称、品牌词归类、品牌词出现详情、收录效果状态、文章链接列表、备注（异常都写清楚，如\"未搜索到\"\"该链接无法解析\"）。", "输出：")

add_h("3. 基于真实使用反馈，再迭代一轮")
add_para("工具真用起来后，我主动反问 AI：\"你觉得这个应用还有哪些值得优化的地方？\" 拿到建议后，挑了最影响使用的几项做了第二版。")
add_image(4, "图 4：自己不清楚能优化到什么程度时，反过来让 AI 给优化清单，再决定做哪些。")
add_image(6, "图 5：第二版新增能力一览——汇总看板、一键启动 start.bat、历史记录面板等。")
add_bullet("固定浏览器标识减少差异；触发验证码时自动等待重试，而不是一上来就判失败；新增\"跑 2~3 次取多数决\"的趋势结论，并标注数据新鲜度。", "稳定性增强：")
add_bullet("跑完直接出图表，一眼看清两个品牌谁收录更好、\"好 / 差\"分布。", "汇总看板：")
add_bullet("导出拆成两张表，链接明细表用超链接公式，点一下就能在浏览器打开。", "链接可点击：")
add_bullet("做了 start.bat，双击就能开，不用每次敲命令。", "一键启动：")
add_bullet("每次跑完自动存档，能回看、对比不同时期的收录变化。", "历史记录面板：")

# ---- R ----
add_h("四、成果（Result）")
add_bullet("原本人工逐个搜索、逐篇对比、手动计数的工作，现在上传一张表、点一下就出报告，几十个产品几分钟搞定；", "")
add_bullet("工具直接告诉每个产品\"更适合用哪个品牌继续投内容\"，并标出收录效果好坏，让品牌运营决策从\"凭感觉\"变成\"看数据\"；", "")
add_bullet("对\"未搜索到\"\"链接无法解析\"\"触发验证\"等边界情况都有备注，方便人工复核，不会悄悄出错；", "")
add_bullet("通过\"跑多次取趋势 + 历史对比\"，结果更可信，也能跟踪一个产品收录随时间变好还是变差。", "")
add_para("优化前 vs 优化后对比：")
add_image(5, "图 6：左边是第一次跑出的原始结果表，右边是第二轮增加\"汇总看板\"后的效果——多了图表，结论一眼可见。")

# ---- 亮点 ----
add_h("五、这个作品说明我能做什么（岗位匹配亮点）")
add_bullet("把\"百度收录对比\"这个模糊的运营问题，拆解成清晰、可落地的功能清单。", "能发现业务痛点并定义需求：")
add_bullet("非技术背景也能借助多个 AI（需求梳理 + 代码实现）完成从 0 到 1 的搭建，并持续迭代。", "会用 AI 工具把想法变成产品：")
add_bullet("主动关注结果可信度（两次结果不一致问题）、边界情况处理、人工复核机制。", "有数据严谨性意识：")
add_bullet("基于真实使用反馈（结果不稳、链接点不开）持续改进，而不是做完就扔。", "有产品迭代思维：")
add_bullet("自己不懂技术细节时，懂得先让 AI 帮梳理需求、再让 AI 给优化清单，把\"不会写代码\"变成\"会指挥 AI\"。", "会借力、会提问：")

# ---- 附录表 ----
add_h("附：面试可展开的技术要点（备查）")
table = doc.add_table(rows=1, cols=2)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
set_cjk(hdr[0].paragraphs[0].add_run("模块")).bold = True
set_cjk(hdr[1].paragraphs[0].add_run("做法")).bold = True
rows = [
    ("反爬", "Playwright 真实浏览器 + 固定 UA + 抹 webdriver 痕迹"),
    ("计数规则", "单篇文章对两个品牌各只计一次，总提及 = 提 A 篇数 + 提 B 篇数（跨品牌相加）"),
    ("品牌词判断", "默认本地精确匹配（离线），可选大模型 API（OpenAI 格式）更精准"),
    ("边界处理", "重复产品、未搜索到、链接无法解析、触发验证码 均单独备注"),
    ("输出", "6 列 Excel + 两张表（结果 / 链接明细）+ 汇总图表 + 历史存档"),
]
for mod, how in rows:
    cells = table.add_row().cells
    set_cjk(cells[0].paragraphs[0].add_run(mod))
    set_cjk(cells[1].paragraphs[0].add_run(how))

out = f"{BASE}\\百度收录检测工具_STAR.docx"
doc.save(out)
print("SAVED:", out)
